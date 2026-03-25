"""
Word 导出服务。

这个模块把统一财务报告导出为 `.docx` 文件，
用于正式汇报、归档或发送给管理层。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from sqlalchemy.orm import Session

from core.config import settings
from schemas.export import ExportFileResponse
from schemas.report import FinancialReportResponse
from services.reporting.report_chart_service import generate_financial_charts
from services.reporting.report_service import generate_financial_report


def _ensure_exports_dir() -> Path:
    """
    确保导出目录存在。
    """
    export_dir = Path(settings.EXPORTS_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)
    return export_dir


def export_financial_word_report(
    db: Session,
    memory_db: Session,
    period: str,
    report: FinancialReportResponse | None = None,
    model: str | None = None,
) -> ExportFileResponse:
    """
    生成指定期间的 Word 财务报告。

    如果上游已经生成过统一财务报告，可以直接把 report 传进来，
    避免这里再次调用模型和重复查询数据库。
    """
    if report is None:
        report = generate_financial_report(db, memory_db, period, model=model)

    chart_paths = generate_financial_charts(period, report.objective_summary)

    document = Document()
    document.styles["Normal"].font.name = "Arial"
    document.styles["Heading 1"].font.name = "Arial"
    document.styles["Heading 2"].font.name = "Arial"

    # 明确设置页边距，保证导出的版式在常见阅读器中更稳定。
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = document.add_heading(f"{period} 财务报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("一、关键指标", level=2)
    metrics_table = document.add_table(rows=1, cols=2)
    metrics_table.style = "Table Grid"
    header_cells = metrics_table.rows[0].cells
    header_cells[0].text = "指标"
    header_cells[1].text = "数值"
    for metric in report.metrics:
        row = metrics_table.add_row().cells
        row[0].text = metric.label
        row[1].text = str(metric.value)

    document.add_heading("二、客观财务数据", level=2)
    summary = report.objective_summary
    document.add_paragraph(f"凭证数量：{summary.voucher_count}")
    document.add_paragraph(f"总收入：{summary.total_revenue}")
    document.add_paragraph(f"总支出：{summary.total_expense}")
    document.add_paragraph(f"净利润：{summary.net_profit}")
    document.add_paragraph(f"待收款总额：{summary.pending_receivables}")
    document.add_paragraph(f"待付款总额：{summary.pending_payables}")

    document.add_heading("三、图表概览", level=2)
    if chart_paths:
        for chart_path in chart_paths:
            document.add_picture(chart_path, width=Inches(6.5))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        document.add_paragraph("当前期间暂无可生成的图表。")

    document.add_heading("四、财务概览", level=2)
    document.add_paragraph(report.narrative.executive_summary)

    document.add_heading("五、财务分析", level=2)
    document.add_paragraph(report.narrative.analysis)

    document.add_heading("六、风险与异常", level=2)
    if report.narrative.risks:
        for risk in report.narrative.risks:
            document.add_paragraph(risk, style="List Bullet")
    else:
        document.add_paragraph("暂无明显风险提示。")

    document.add_heading("七、建议动作", level=2)
    if report.narrative.recommendations:
        for item in report.narrative.recommendations:
            document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph("暂无建议动作。")

    export_dir = _ensure_exports_dir()
    filename = f"fina-report-{period}.docx"
    output_path = export_dir / filename
    document.save(output_path)

    return ExportFileResponse(
        file_type="docx",
        period=period,
        filename=filename,
        path=str(output_path.resolve()),
    )
